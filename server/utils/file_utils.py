import os
from io import BytesIO
import PyPDF2
from docx import Document as DocxDocument
import pdfplumber
import re
import logging
from datetime import datetime
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = os.getenv('ALLOWED_EXTENSIONS', 'pdf,docx').split(',')
MAX_SECTION_CHECK = int(os.getenv('MAX_SECTION_CHECK', 100))

class FileProcessingError(Exception):
    """Custom exception for file processing errors"""
    pass

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_stream):
    try:
        pdf_reader = PyPDF2.PdfReader(file_stream)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
        return text
    except PyPDF2.errors.PdfReadError as e:
        logger.error(f"PDF read error: {str(e)}")
        raise FileProcessingError(f"Failed to read PDF: {str(e)}")
    except ValueError as e:
        logger.error(f"Invalid PDF data: {str(e)}")
        raise FileProcessingError(f"Invalid PDF data: {str(e)}")

def extract_text_from_docx(file_stream):
    try:
        doc = DocxDocument(file_stream)
        return "\n".join([para.text for para in doc.paragraphs])
    except PackageNotFoundError as e:
        logger.error(f"DOCX file not found or corrupted: {str(e)}")
        raise FileProcessingError(f"DOCX file corrupted: {str(e)}")
    except ValueError as e:
        logger.error(f"Invalid DOCX data: {str(e)}")
        raise FileProcessingError(f"Invalid DOCX data: {str(e)}")

def extract_docx_metadata(file_path: str) -> dict:
    metadata = {
        "title": os.path.basename(file_path),
        "author": "Unknown",
        "keywords": "",
        "subject": "",
        "is_research": False,
        "total_pages": 0,
        "sections": []
    }
    try:
        doc = DocxDocument(file_path)
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
        if doc.core_properties.keywords:
            metadata["keywords"] = doc.core_properties.keywords
        if doc.core_properties.subject:
            metadata["subject"] = doc.core_properties.subject
        total_paragraphs = len(doc.paragraphs)
        metadata["total_pages"] = max(1, total_paragraphs // 30)
        text = "\n".join([para.text for para in doc.paragraphs[:MAX_SECTION_CHECK]])
        metadata["is_research"] = any(
            re.search(pattern, text, re.IGNORECASE)
            for pattern in [r'abstract', r'introduction', r'methodology', r'references']
        )
        for para in doc.paragraphs[:MAX_SECTION_CHECK]:
            if para.style.name.lower().startswith('heading'):
                metadata["sections"].append(para.text.strip())
    except PackageNotFoundError as e:
        logger.error(f"DOCX metadata extraction error: {str(e)}")
        raise FileProcessingError(f"DOCX file issue: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected DOCX metadata error: {str(e)}")
        raise FileProcessingError(f"Unexpected error in DOCX metadata: {str(e)}")
    return metadata

def extract_pdf_metadata(file_path: str) -> dict:
    metadata = {
        "title": os.path.basename(file_path),
        "author": "Unknown",
        "keywords": "",
        "subject": "",
        "is_research": False,
        "image_count": 0,
        "figure_count": 0,
        "table_count": 0,
        "total_pages": 0,
        "sections": []
    }
    try:
        with pdfplumber.open(file_path) as pdf:
            metadata["total_pages"] = len(pdf.pages)
            if hasattr(pdf, 'metadata'):
                pdf_meta = pdf.metadata or {}
                metadata.update({
                    "title": pdf_meta.get('Title', metadata['title']),
                    "author": pdf_meta.get('Author', metadata['author']),
                    "keywords": pdf_meta.get('Keywords', metadata['keywords']),
                    "subject": pdf_meta.get('Subject', metadata['subject'])
                })
            first_page_text = pdf.pages[0].extract_text() or ""
            metadata["is_research"] = any(
                re.search(pattern, first_page_text, re.IGNORECASE)
                for pattern in [r'abstract', r'introduction', r'methodology', r'references']
            )
            for page in pdf.pages[:MAX_SECTION_CHECK]:
                text = page.extract_text() or ""
                metadata["figure_count"] += len(re.findall(r'(?:Figure|Fig\.?)\s*\d+', text, re.IGNORECASE))
                metadata["table_count"] += len(re.findall(r'(?:Table|Tab\.?)\s*\d+', text, re.IGNORECASE))
                metadata["image_count"] += len(page.images)
                if page.page_number == 1:
                    section_matches = re.findall(r'^(?:[1-9]\.\s+)?([A-Z][A-Za-z\s]+?)\s*$', text, re.MULTILINE)
                    metadata["sections"] = [s.strip() for s in section_matches if len(s.strip()) > 5]
    except Exception as e:
        logger.error(f"PDF metadata extraction error: {str(e)}")
        raise FileProcessingError(f"Failed to extract PDF metadata: {str(e)}")
    return metadata

def extract_metadata(file_path: str) -> dict:
    if file_path.endswith('.pdf'):
        return extract_pdf_metadata(file_path)
    elif file_path.endswith('.docx'):
        return extract_docx_metadata(file_path)
    return {}